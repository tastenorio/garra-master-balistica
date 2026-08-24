import logging
import warnings
import os

# Tentativas de silenciar avisos nativos do Google
os.environ["GRPC_VERBOSITY"] = "ERROR"
os.environ["GLOG_minloglevel"] = "2"
logging.getLogger("google").setLevel(logging.ERROR)
warnings.filterwarnings("ignore")

import fitz  # PyMuPDF
import json
import re
import io
import ollama
import chromadb
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google import genai
from google.genai import types
from PIL import Image

# --- 1. CONFIGURAÇÃO DA API GEMINI E BANCO DE DADOS (RAG) ---
CHAVE_API_GEMINI = "cole-aqui-sua-chave-api"
cliente_gemini = genai.Client(api_key=CHAVE_API_GEMINI)

# Desativando Filtros de Segurança
CONFIG_SEGURANCA = [
    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE")
]

try:
    cliente_db = chromadb.PersistentClient(path="./banco_balistica")
    colecao_manuais = cliente_db.get_collection(name="manuais_armas")
    USA_RAG = True
except Exception as e:
    print(f"Aviso: Banco de dados não encontrado ou vazio.")
    USA_RAG = False

# --- 2. PROMPTS INTELIGENTES ---
PROMPT_CLASSIFICADOR = """Analise esta imagem e classifique o seu conteúdo.
Responda APENAS com a palavra "DOCUMENTO" ou "EVIDENCIA".
- DOCUMENTO: Se for um ofício, requisição policial, folha de papel com texto impresso ou formulário.
- EVIDENCIA: Se for arma, munição, estojo, projétil, régua, ou um envelope/invólucro pardo/plástico fechado contendo materiais.
Responda APENAS UMA PALAVRA."""

PROMPT_PDF_TEXTO = """Você é um assistente técnico de Perícia Criminal. Extraia os dados do ofício policial fornecido abaixo no formato JSON.

REGRAS DE FORMATAÇÃO OBRIGATÓRIAS:
1. Formate TODOS os nomes (pessoas, delegacias, autoridades) em 'Title Case' (Iniciais Maiúsculas).
2. Escreva as datas RIGOROSAMENTE no formato "DD de [mês por extenso] de YYYY" (ex: "11 de janeiro de 1985").
3. Se não encontrar uma informação, preencha EXATAMENTE com a palavra "[PREENCHER]".
4. ABSOLUTAMENTE PROIBIDO: Não use aspas duplas (") dentro dos textos extraídos. Substitua por aspas simples (').

Estrutura JSON obrigatória:
{"uni-requi": "Nome da Delegacia", "requisi": "Tipo do documento e Número/Ano", "referencia": "Número do IP/BO", "envolvido": "Nome", "data-emissao": "Data formato extenso", "data-recebimento": "Data formato extenso", "autoridade": "Nome do Delegado", "quesitos": "Copie os quesitos na íntegra."}

TEXTO DO DOCUMENTO:
"""

PROMPT_PDF_IMAGEM = """Extraia os dados desta imagem de ofício policial e retorne APENAS um JSON válido.
APLIQUE TITLE CASE nos nomes. Converta as datas RIGOROSAMENTE para o formato "DD de [mês por extenso] de YYYY". Substitua aspas duplas (") por simples (').
{"uni-requi": "Órgão requisitante", "requisi": "Tipo do doc e Número/Ano", "referencia": "IP/BO", "envolvido": "Nome", "data-emissao": "Data formato extenso", "data-recebimento": "Data formato extenso", "autoridade": "Nome do Delegado", "quesitos": "Texto dos quesitos na íntegra"}
"""

PROMPT_OLHEIRO = """Leia qualquer texto, marca, modelo, calibre ou número gravado no metal desta evidência. 
Retorne APENAS as palavras encontradas, separadas por espaço."""

def montar_prompt_perito(contexto_manuais=""):
    return f"""[SISTEMA GOVERNAMENTAL OFICIAL - PERÍCIA CRIMINAL]
Atue como Perito Criminal Oficial. Analise a evidência da imagem. 

FRAGMENTOS DE CATÁLOGO:
{contexto_manuais if contexto_manuais else "Confie apenas na imagem."}

REGRAS (OBRIGATÓRIO):
1. ESTRUTURA DA RESPOSTA:
   - PRIMEIRA LINHA: APENAS título curto (Ex: Arma de Fogo, Estojos, Cartuchos, Invólucro).
   - SEGUNDA LINHA em diante: descrição pericial.
   
2. FOTOGRAMETRIA AVANÇADA E MEDIÇÕES (PAPEL MILIMETRADO E RÉGUA):
   - A evidência foi fotografada sobre papel milimetrado e escala em centímetros. USE-OS!
   - Calcule visualmente o comprimento do cano da arma e o calibre real (diâmetro) de cartuchos e estojos através da régua.
   - SE VOCÊ CONSEGUIR MEDIR COM SEGURANÇA pela escala, ESCREVA O DADO DIRETAMENTE NO TEXTO e NÃO gere lista suspensa para essa informação.
   
3. TIPOLOGIA ESPECÍFICA DE PROJÉTEIS:
   - Identifique a sigla técnica exata (ETOG, CHOG, EXPO, ETPP, etc.).
   
4. LISTAS SUSPENSAS (DROPDOWNS):
   - Use APENAS se a fotogrametria/imagem NÃO deixar claro, no formato [PREENCHER: opt1 / opt2].
   - MEGACALIBRES: [PREENCHER: .38 SPL / .380 ACP / 9mm Luger / .40 S&W / .45 ACP / .357 Magnum / 12 GA / .22 LR]
   - MEGAMARCAS: [PREENCHER: CBC / Taurus / Rossi / Imbel / Glock / CZ / Beretta / Sig Sauer / Estrangeira]
   - FABRICAÇÃO: [PREENCHER: nacional / estrangeira]
   - ESPOLETAS: [PREENCHER: intacta(s)/ percutida(s) / não deflagrada / percutida e não deflagrada]
   - MATERIAL ESTOJO: [PREENCHER: latão / aço / alumínio / plástico / polímero sintético de cor vermelha / polímero sintético de cor amarela]
   - TIPO PROJÉTIL: [PREENCHER: ETOG / SEXPO /CHOG / EXPO / ETPP / semi-canto-vivo / chumbo nu]

5. SE FOR UM INVÓLUCRO/EMBALAGEM DE RECEBIMENTO:
   - A PRIMEIRA LINHA deve ser EXATAMENTE a palavra: Invólucro
   - A SEGUNDA LINHA deve ser UMA ÚNICA FRASE RESUMIDA e EM TEXTO CORRIDO descrevendo o material, lacre e texto manuscrito.
   - OBRIGATÓRIO: Inicie a descrição com LETRA MINÚSCULA (ela será inserida no meio de uma frase do laudo).
   - Exemplo 1: envelope de papel pardo, fechado por grampos metálicos e sem inscrições manuscritas visíveis
   - Exemplo 2: embalagem plástica reaproveitada, fechada com clipe metálico e com inscrições 'IP 04/24'

6. SE FOR ARMA DE FOGO:
   - Texto contínuo (prosa). Inicie: "Trata-se de..." 

7. SE FOR ESTOJO(S) OU CARTUCHO(S):
   - Tópicos verticais com UMA QUEBRA DE LINHA (Enter) após o ponto e vírgula.
   - Natureza: [PREENCHER: cartucho(s) de arma de fogo / estojo(s) de cartucho de arma de fogo];
   - Quantidade: [PREENCHER: 01 / 02 / 03 / 04 / 05 / 06]; Fabricação: [Usar FABRICAÇÃO];
   - Código de lote: [PREENCHER: ausente / ilegível];
   - Marca: [Usar MEGAMARCAS];
   - Calibre nominal: [Usar FOTOGRAMETRIA ou MEGACALIBRES];
   - Estado da(s) espoleta(s): [Usar ESPOLETAS];
   - Construção/ material do(s) estojo(s): [Usar MATERIAL ESTOJO];
   - Projétil(eis): [Se cartucho, usar TIPOLOGIA ESPECÍFICA].

8. SE FOR PROJÉTIL(EIS):
   - Inicie: "Trata-se de [PREENCHER: 01 / 02 / 03] projétil(eis)...".
   - Tabela: Tipo | Massa | Calibre (Via Fotogrametria) | Deformações normais (raias) | Deformações acidentais.

RETORNE APENAS TÍTULO E DESCRIÇÃO.
"""

# --- 3. FUNÇÕES DE PROCESSAMENTO ---
def classificar_imagem(caminho_imagem):
    try:
        img = Image.open(caminho_imagem)
        config_classificador = types.GenerateContentConfig(temperature=0.0)
        resposta = cliente_gemini.models.generate_content(
            model='gemini-3.5-flash',
            contents=[PROMPT_CLASSIFICADOR, img],
            config=config_classificador
        )
        resultado = resposta.text.strip().upper() if resposta and resposta.text else "EVIDENCIA"
        return "DOCUMENTO" if "DOCUMENTO" in resultado else "EVIDENCIA"
    except Exception as e:
        return "EVIDENCIA"

def extrair_dados_requisicao(caminho):
    try:
        config_json = types.GenerateContentConfig(
            response_mime_type="application/json", 
            temperature=0.0,
            safety_settings=CONFIG_SEGURANCA
        )
        
        if caminho.lower().endswith('.pdf'):
            doc = fitz.open(caminho)
            texto_completo = "\n".join([pagina.get_text("text") for pagina in doc])
            
            if len(texto_completo.strip()) > 50:
                print("      [Nuvem]: Extraindo dados textuais do Ofício PDF (Gemini 3.5)...")
                resposta = cliente_gemini.models.generate_content(
                    model='gemini-3.5-flash',
                    contents=PROMPT_PDF_TEXTO + texto_completo,
                    config=config_json
                )
                return resposta.text.strip() if resposta and resposta.text else ""
                
            print("      [Nuvem]: Ofício PDF escaneado detectado. Analisando imagem (Gemini 3.5)...")
            pix = doc.load_page(0).get_pixmap(matrix=fitz.Matrix(2, 2))
            imagem_req = Image.open(io.BytesIO(pix.tobytes("png")))
        else:
            print("      [Nuvem]: Ofício em formato de Imagem detectado. Analisando (Gemini 3.5)...")
            imagem_req = Image.open(caminho)
            
        resposta = cliente_gemini.models.generate_content(
            model='gemini-3.5-flash',
            contents=[PROMPT_PDF_IMAGEM, imagem_req],
            config=config_json
        )
        return resposta.text.strip() if resposta and resposta.text else ""
    except Exception as e:
        print(f"      [Erro na Requisição]: {e}")
        return ""

def processar_foto_com_rag(caminho):
    try:
        imagem_evidencia = Image.open(caminho)
        contexto_manuais = ""
        config_foto = types.GenerateContentConfig(
            temperature=0.2,
            safety_settings=CONFIG_SEGURANCA
        )
        
        if USA_RAG:
            print("      [Olheiro]: Escaneando imagem...")
            resp_olheiro = cliente_gemini.models.generate_content(
                model='gemini-3.5-flash',
                contents=[PROMPT_OLHEIRO, imagem_evidencia],
                config=config_foto
            )
            termos_busca = resp_olheiro.text.strip() if resp_olheiro and resp_olheiro.text else ""
            
            if termos_busca:
                try:
                    resp_embed = ollama.embeddings(model="nomic-embed-text", prompt=termos_busca)
                    resultados = colecao_manuais.query(query_embeddings=[resp_embed["embedding"]], n_results=2)
                    if resultados['documents'] and resultados['documents'][0]:
                        contexto_manuais = "\n---\n".join(resultados['documents'][0])
                except:
                    pass 

        print("      [Perito]: Analisando Fotogrametria e descrevendo material...")
        prompt_final = montar_prompt_perito(contexto_manuais)
        resposta = cliente_gemini.models.generate_content(
            model='gemini-3.5-flash',
            contents=[prompt_final, imagem_evidencia],
            config=config_foto
        )
        
        if not resposta or not resposta.text:
            return "[PREENCHER]\n[Falha visual: Filtro de segurança da API ativado ou erro de conexão]"
            
        return resposta.text.strip()
    except Exception as e:
        print(f"      [Erro na Foto]: {e}")
        return "[PREENCHER]\n[Falha na descrição visual desta evidência]"

# --- FORMATAÇÃO DE QUESITOS ---
def formatar_quesitos(texto_bruto):
    if not texto_bruto or texto_bruto.strip().lower() in ["", "[preencher]", "preencher"]:
        return "Não foram formulados quesitos específicos por parte da autoridade policial."
    
    texto_bruto = re.sub(r'\?', '?\n', texto_bruto)
    texto_bruto = re.sub(r'\*', '\n*', texto_bruto)
    
    linhas = texto_bruto.split('\n')
    quesitos_formatados = []
    contador = 1
    
    for linha in linhas:
        l_limpa = linha.strip()
        l_limpa = l_limpa.replace('*', '') 
        l_limpa = re.sub(r'^[\-\•\>]\s*', '', l_limpa)
        l_limpa = re.sub(r'^\d+[\.\)\-]\s*', '', l_limpa)
        
        if len(l_limpa) > 3:
            quesitos_formatados.append(f"{contador}. {l_limpa}")
            contador += 1
            
    return "\n\n".join(quesitos_formatados) if quesitos_formatados else "[PREENCHER]"

# --- 4. MOTOR PRINCIPAL E INJEÇÃO XML SEGURO ---
def substituir_no_paragrafo(p, tag, valor):
    if tag not in p.text: return
    substituido = False
    for run in p.runs:
        if tag in run.text:
            run.text = run.text.replace(tag, str(valor))
            substituido = True
    if not substituido:
        run_base = None
        for run in p.runs:
            if run.text.strip():
                run_base = run
                break
        font_name = run_base.font.name if run_base else None
        font_size = run_base.font.size if run_base else None
        novo_texto = p.text.replace(tag, str(valor))
        p.clear()
        partes = novo_texto.split('\n')
        for idx, parte in enumerate(partes):
            if parte or len(partes) == 1:
                novo_run = p.add_run(parte)
                if font_name: novo_run.font.name = font_name
                if font_size: novo_run.font.size = font_size
            if idx < len(partes) - 1: p.add_run().add_break()

def substituir_mantendo_formatacao(doc, substituicoes):
    for paragrafo in doc.paragraphs:
        for tag, valor in substituicoes.items():
            substituir_no_paragrafo(paragrafo, tag, valor)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for tag, valor in substituicoes.items():
                        substituir_no_paragrafo(paragrafo, tag, valor)

def processar_dropdowns_xml(doc):
    regex_geral = re.compile(r'(\[PREENCHER[^\]]*\])', re.IGNORECASE)
    
    def formatar_paragrafo(p):
        if not ('[PREENCHER' in p.text.upper()): return
            
        run_base = None
        for r in p.runs:
            if r.text.strip():
                run_base = r
                break
                
        nome_fonte = run_base.font.name if run_base else None
        tamanho_fonte = run_base.font.size if run_base else None
        
        partes = regex_geral.split(p.text)
        p.clear() 
        
        for parte in partes:
            if not parte: continue
            
            if parte.upper().startswith('[PREENCHER:'):
                opcoes_str = parte[11:-1]
                opcoes = [o.strip() for o in opcoes_str.split('/') if o.strip()]
                if not opcoes: opcoes = ["Selecione"]
                
                sdt = OxmlElement('w:sdt')
                sdtPr = OxmlElement('w:sdtPr')
                dropDownList = OxmlElement('w:dropDownList')
                
                for opt in opcoes:
                    listItem = OxmlElement('w:listItem')
                    listItem.set(qn('w:displayText'), opt)
                    listItem.set(qn('w:value'), opt)
                    dropDownList.append(listItem)
                    
                sdtPr.append(dropDownList)
                sdt.append(sdtPr)
                
                sdtContent = OxmlElement('w:sdtContent')
                r = OxmlElement('w:r')
                
                rPr = OxmlElement('w:rPr')
                color = OxmlElement('w:color')
                color.set(qn('w:val'), 'FF0000')
                b = OxmlElement('w:b')
                
                if nome_fonte:
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), nome_fonte)
                    rFonts.set(qn('w:hAnsi'), nome_fonte)
                    rPr.append(rFonts)
                
                if tamanho_fonte:
                    try:
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), str(int(tamanho_fonte.pt * 2)))
                        rPr.append(sz)
                    except: pass
                    
                rPr.append(color)
                rPr.append(b)
                r.append(rPr)
                
                t = OxmlElement('w:t')
                t.text = opcoes[0]
                r.append(t)
                
                sdtContent.append(r)
                sdt.append(sdtContent)
                p._p.append(sdt)
                
            elif parte.upper() == '[PREENCHER]':
                run_alerta = p.add_run('[PREENCHER]')
                if nome_fonte: run_alerta.font.name = nome_fonte
                if tamanho_fonte: run_alerta.font.size = tamanho_fonte
                run_alerta.font.color.rgb = RGBColor(255, 0, 0)
                run_alerta.font.bold = True
            else:
                run_normal = p.add_run(parte)
                if nome_fonte: run_normal.font.name = nome_fonte
                if tamanho_fonte: run_normal.font.size = tamanho_fonte

    for paragrafo in doc.paragraphs:
        formatar_paragrafo(paragrafo)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    formatar_paragrafo(paragrafo)

if __name__ == "__main__":
    print("Iniciando Garra Master v11.4 - Cadeia de Custódia Integrada...\n")
    
    pasta_destino_laudos = r"C:\Users\balis\OneDrive\ICRIM\LAUDOS\Balistica\CARACTERIZACAO\Natureza-e-eficiencia\2026\08-2026\laudos"
    os.makedirs(pasta_destino_laudos, exist_ok=True)
    
    template_caminho = "template_laudo.docx"
    pastas_ignoradas = ['venv', 'manuais', 'banco_balistica', '__pycache__', 'laudos']
    pastas_casos = [d for d in os.listdir('.') if os.path.isdir(d) and d not in pastas_ignoradas and not d.startswith('.')]

    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    hoje = datetime.now()
    data_hoje_extenso = f"{hoje.day:02d} de {meses[hoje.month - 1]} de {hoje.year} às {hoje.strftime('%Hh%M')}"

    for pasta in pastas_casos:
        print(f"\n=======================================")
        caminho_pasta = os.path.join('.', pasta)
        arquivo_marcador = os.path.join(caminho_pasta, ".processado")
        
        if os.path.exists(arquivo_marcador):
            print(f"⏩ Pulando caso '{pasta}': Já processado anteriormente.")
            continue
            
        print(f"📂 PROCESSANDO CASO: {pasta}")
        
        arquivos = os.listdir(caminho_pasta)
        
        pdfs = [f for f in arquivos if f.lower().endswith('.pdf')]
        fotos = [f for f in arquivos if f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))]
        
        arquivo_req = None
        fotos_evidencia = []

        if pdfs:
            arquivo_req = pdfs[0]
            fotos_evidencia = fotos
        elif fotos:
            print(f"   [Sistema]: Iniciando Pré-Visão IA para separar Ofício das Evidências...")
            for foto in fotos:
                caminho_foto = os.path.join(caminho_pasta, foto)
                
                if not arquivo_req:
                    classe = classificar_imagem(caminho_foto)
                    if classe == "DOCUMENTO":
                        arquivo_req = foto
                        print(f"      -> {foto}: [DOCUMENTO DETECTADO]")
                    else:
                        fotos_evidencia.append(foto)
                        print(f"      -> {foto}: [EVIDÊNCIA]")
                else:
                    fotos_evidencia.append(foto)
                    print(f"      -> {foto}: [EVIDÊNCIA]")
                    
        if not arquivo_req:
             print(f"   [!] Ignorando pasta '{pasta}': Nenhum PDF ou imagem de Ofício impresso foi detectado pela IA.")
             continue
             
        if not fotos_evidencia:
             print(f"   [!] Ignorando pasta '{pasta}': Requisição encontrada ({arquivo_req}), mas faltam fotos das evidências.")
             continue

        print(f"\n   📄 Lendo Requisição Policial ({arquivo_req})...")
        resposta_pdf = extrair_dados_requisicao(os.path.join(caminho_pasta, arquivo_req))
        
        try:
            if not resposta_pdf:
                dados_cartorio = {}
            else:
                idx_inicio = resposta_pdf.find('{')
                if idx_inicio != -1:
                    texto_para_decodificar = resposta_pdf[idx_inicio:]
                    decodificador = json.JSONDecoder()
                    dados_cartorio, _ = decodificador.raw_decode(texto_para_decodificar)
                else:
                    dados_cartorio = {}
        except Exception as e:
            print(f"      [Erro de Conversão JSON]: Falha na decodificação bruta.")
            dados_cartorio = {}

        quesitos_extraidos = dados_cartorio.get("quesitos", "[PREENCHER]")
        if isinstance(quesitos_extraidos, list):
            quesitos_extraidos = "\n".join(str(q) for q in quesitos_extraidos)
        elif quesitos_extraidos is None:
            quesitos_extraidos = "[PREENCHER]"
        else:
            quesitos_extraidos = str(quesitos_extraidos)

        quesitos_formatados = formatar_quesitos(quesitos_extraidos)

        descricoes_involucros = []
        descricoes_finais = []
        titulos_evidencias = []
        
        for i, foto in enumerate(fotos_evidencia):
            print(f"   📸 Analisando material {i+1} ({foto})...")
            texto_evidencia = processar_foto_com_rag(os.path.join(caminho_pasta, foto))
            
            linhas = texto_evidencia.split('\n', 1)
            titulo = linhas[0].replace('*', '').strip()
            descricao = linhas[1].strip() if len(linhas) > 1 else ""
            
            titulos_evidencias.append(titulo)
            
            # --- LÓGICA DE INJEÇÃO DO INVÓLUCRO (V11.4) ---
            if titulo.upper() == "INVÓLUCRO":
                # Força a primeira letra ser minúscula
                if descricao:
                    descricao = descricao[0].lower() + descricao[1:]
                descricoes_involucros.append(descricao)
            else:
                letra = chr(97 + len(descricoes_finais))
                descricoes_finais.append(f"{letra}) {titulo}:\n{descricao}")
            
        # Tratativa elegante para múltiplos invólucros na mesma frase
        if descricoes_involucros:
            if len(descricoes_involucros) > 1:
                bloco_involucros = ", ".join(descricoes_involucros[:-1]) + " e " + descricoes_involucros[-1]
            else:
                bloco_involucros = descricoes_involucros[0]
        else:
            bloco_involucros = "ausência de registro fotográfico de invólucro ou embalagem de recebimento"
            
        bloco_descritivo_completo = "\n\n".join(descricoes_finais) if descricoes_finais else "[PREENCHER: Nenhuma evidência balística processada]"

        print(f"   📝 Montando Laudo e injetando fotos (Largura 11cm)...")
        substituicoes = {
            "{{DATA_PERICIA}}": data_hoje_extenso,
            "{{ORGAO_REQUISITANTE}}": dados_cartorio.get("uni-requi", "[PREENCHER]"),
            "{{AUTORIDADE}}": dados_cartorio.get("autoridade", "[PREENCHER]"),
            "{{NUM_REQUISICAO}}": dados_cartorio.get("requisi", "[PREENCHER]"),
            "{{DATA_EMISSAO}}": dados_cartorio.get("data-emissao", "[PREENCHER]"),
            "{{DATA_RECEBIMENTO}}": dados_cartorio.get("data-recebimento", "[PREENCHER]"),
            "{{REFERENCIA}}": dados_cartorio.get("referencia", "[PREENCHER]"),
            "{{ENVOLVIDO}}": dados_cartorio.get("envolvido", "[PREENCHER]"),
            "{{QUESITOS}}": quesitos_formatados,
            "{{DESC_INVOLUCRO}}": bloco_involucros,
            "{{DESC_MATERIAL}}": bloco_descritivo_completo
        }
        
        nome_laudo_seguro = re.sub(r'[\\/*?:"<>|]', '-', dados_cartorio.get('requisi', 'Novo'))
        nome_laudo = f"Laudo_Preenchido_{nome_laudo_seguro}.docx"
        caminho_saida = os.path.join(pasta_destino_laudos, nome_laudo)
        
        try:
            doc = Document(template_caminho)
            substituir_mantendo_formatacao(doc, substituicoes)
            
            tag_encontrada = False
            for p in doc.paragraphs:
                if "{{FOTOS}}" in p.text:
                    tag_encontrada = True
                    p.text = p.text.replace("{{FOTOS}}", "")
                    
                    for i, foto in enumerate(fotos_evidencia):
                        caminho_foto = os.path.join(caminho_pasta, foto)
                        
                        p_img = p.insert_paragraph_before()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        run_img.add_picture(caminho_foto, width=Cm(11.0))
                        
                        legenda = f"Fotografia {i+1}: {titulos_evidencias[i].lower()}."
                        p_leg = p.insert_paragraph_before(legenda)
                        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.insert_paragraph_before("")
                    break
            
            if not tag_encontrada:
                print("      [Aviso]: Tag {{FOTOS}} não encontrada no template.")
                doc.add_page_break() 
                titulo_apendice = doc.add_heading('VI. APÊNDICE FOTOGRÁFICO', level=1)
                titulo_apendice.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, foto in enumerate(fotos_evidencia):
                    caminho_foto = os.path.join(caminho_pasta, foto)
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(caminho_foto, width=Cm(11.0))
                    
                    legenda = f"Fotografia {i+1}: {titulos_evidencias[i].lower()}."
                    p_leg = doc.add_paragraph(legenda)
                    p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph() 
            
            print("      [Sistema]: Injetando dropdowns XML interativos...")
            processar_dropdowns_xml(doc)
            
            doc.save(caminho_saida)
            
            with open(arquivo_marcador, 'w', encoding='utf-8') as f:
                f.write(f"Processado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\nLaudo gerado: {nome_laudo}")
                
            print(f"   ✅ SUCESSO! Laudo salvo em: {caminho_saida}")
            
        except Exception as e:
            print(f"   [ERRO] Falha ao salvar o documento: {e}")
            
    print("\n🏁 Processamento em lote concluído!")
